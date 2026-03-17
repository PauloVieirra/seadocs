"""API FastAPI do serviço RAG."""
import json
import os
import re
from contextlib import asynccontextmanager
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from supabase import create_client, Client

from chroma_manager import get_or_create_collection, get_or_create_client
from document_processor import (
    index_document,
    index_positive_feedback as _index_positive_feedback,
    delete_positive_feedback as _delete_positive_feedback,
    delete_document_from_chroma,
    compute_file_hash,
)
from config import CHROMA_PERSIST_DIR, COLLECTION_NAME

# Diretório de modelos de exemplo (relativo à raiz do projeto)
MODELOS_DIR = Path(__file__).parent.parent / "public" / "modelos"
# Arquivo central de regras constitucionais
REGRAS_PATH = Path(__file__).parent.parent / "Spec" / "REGRAS_CONSTITUCIONAIS.md"

# Fallback se o arquivo não existir (ex.: deploy sem Spec/)
_CONSTITUTIONAL_RULES_FALLBACK = (
    "REGRAS OBRIGATÓRIAS (violação invalida o documento):\n"
    "1. NÃO transcreva, cite ou reproduza trechos literais dos documentos da base. Sintetize.\n"
    "2. NÃO replique partes de conversas extraídas da base (chats, mensagens, diálogos). "
    "INTERPRETE a conversa e crie texto formal que represente o que precisa ser feito — nunca copie diálogos literais.\n"
    "3. NÃO inclua raciocínio, explicações de processo ou planos de ação da IA no texto gerado.\n"
    "   Proibido: frases como 'Vou estruturar...', 'Analisando os dados...', 'Passo 1:...'.\n"
    "4. NÃO mencione nomes de pessoas físicas (autores, participantes, responsáveis, stakeholders).\n"
    "   Use papéis genéricos: [Responsável], [Gestor do Projeto], [Equipe Técnica].\n"
    "5. O documento deve conter APENAS o conteúdo final esperado para a seção.\n"
    "6. O Spec é APENAS regras de estrutura/formato. NUNCA copie, cite ou inclua texto do Spec no documento. "
    "O conteúdo vem SOMENTE dos dados do projeto (base de conhecimento).\n"
    "7. NÃO INVENTE conteúdo. Use APENAS informações presentes nos [DADOS DO PROJETO] recuperados da base. "
    "Se não houver informação para a seção, escreva de forma breve que não há dados disponíveis ou deixe em branco.\n"
    "8. As seções são PARTES de um único documento, NÃO documentos separados. A partir da seção 2, é PROIBIDO repetir introdução, objetivos, visão geral ou contexto.\n"
    "9. NÃO pule para a próxima sessão deixando texto incompleto. Cada seção deve terminar com frases e parágrafos completos. Se precisar mudar de sessão, remova o texto incompleto e complete na sessão adequada (ou crie sessão complementar). Pode encurtar uma sessão e criar outra para o restante — não há limite de sessões.\n"
    "10. Ao finalizar o documento, executar revisão obrigatória. Status na tela: Revisando documento, procurando erros, ajustando sessão X.\n"
    "11. Ignore as tags de markdown no documento gerado; elas são apenas para formatação no editor de texto."
)

# Prompt de correção — acionado quando o usuário escolhe recriar documento (avaliação "não ficou bom" ou "pode melhorar")
_PROMPT_CORRECAO_DOCUMENTO = """
O usuário solicitou RECRIAR o documento porque não ficou bom. Execute as seguintes etapas ao gerar cada seção:

1. ANÁLISE: Identifique seções duplicadas, requisitos repetidos, inconsistências entre níveis (épico, feature, user story, requisito funcional) e problemas de clareza.

2. LIMPEZA: Remova repetições de texto, seções duplicadas, frases redundantes e conteúdo que não adiciona nova informação.

3. REESTRUTURAÇÃO: Organize na estrutura profissional:
   - Visão Geral | Objetivo do Sistema | Escopo (Incluído/Fora)
   - Épicos → Features → User Stories (Como [usuário] Quero [ação] Para que [benefício])
   - Requisitos Funcionais (RF01, RF02...) | Auditoria e Conformidade
   - Cenários Especiais | Benefícios Esperados | Evolução do Produto

4. MELHORIAS: Clareza, precisão técnica, consistência entre seções.

5. NORMALIZAÇÃO: Cada requisito apenas uma vez; features derivam de épicos; user stories de features; RFs objetivos.

6. FORMATAÇÃO: HTML limpo e profissional.

IMPORTANTE: NÃO invente funcionalidades novas. NÃO altere a intenção do sistema. Apenas organize, corrija e consolide.

A seguir está o documento original para revisão:
"""


def _load_constitutional_rules() -> str:
    """Carrega regras constitucionais de Spec/REGRAS_CONSTITUCIONAIS.md."""
    if not REGRAS_PATH.exists():
        return _CONSTITUTIONAL_RULES_FALLBACK
    try:
        text = REGRAS_PATH.read_text(encoding="utf-8")
        for part in text.split("```"):
            if part.strip().startswith("REGRAS OBRIGATÓRIAS"):
                return part.strip()
        return _CONSTITUTIONAL_RULES_FALLBACK
    except Exception:
        return _CONSTITUTIONAL_RULES_FALLBACK


def get_supabase() -> Optional[Client]:
    url = os.getenv("SUPABASE_URL")
    key = os.getenv("SUPABASE_SERVICE_KEY")
    if not url or not key:
        return None
    return create_client(url, key)


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Inicializa ChromaDB na subida do servidor."""
    CHROMA_PERSIST_DIR.mkdir(parents=True, exist_ok=True)
    get_or_create_collection()
    yield
    # cleanup se necessário
    pass


app = FastAPI(title="RAG Service", lifespan=lifespan)

# CORS: permite o frontend chamar a API (qualquer origem em dev)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


class IndexRequest(BaseModel):
    document_id: str
    file_url: str
    file_name: str
    project_id: str
    file_path: Optional[str] = None  # Para download via Supabase Storage (bucket privado)
    reindex: bool = False


class DeleteRequest(BaseModel):
    document_id: str


class PositiveFeedbackSection(BaseModel):
    title: str
    content: str


class IndexPositiveFeedbackRequest(BaseModel):
    project_id: str
    document_id: str
    sections: list[PositiveFeedbackSection]


class DocumentSectionInput(BaseModel):
    id: str
    title: str
    helpText: Optional[str] = None
    repeatable: bool = False
    planningInstruction: Optional[str] = None


class GenerateUnderstandingRequest(BaseModel):
    project_id: str
    document_id: str
    sections: list[DocumentSectionInput]


class GenerateSectionRequest(BaseModel):
    project_id: str
    section_title: str
    help_text: Optional[str] = None
    document_context: Optional[str] = None  # Contexto textual das seções anteriores (legado)
    previous_sections_html: Optional[str] = None  # HTML completo das seções já geradas (partes do mesmo documento)
    section_index: Optional[int] = None  # Índice desta seção (0-based)
    total_sections: Optional[int] = None  # Total de seções do documento
    spec_guidelines: Optional[str] = None  # Diretrizes Spec + Skill (formato: --- SPEC ---\n...\n--- SKILL ---\n...)
    document_type: Optional[str] = None  # Tipo documental do modelo (ex.: Requisito) — a IA usa tipo + spec para montar a estrutura
    creator_feedback: Optional[str] = None  # Avaliação anterior do criador (rating + comment) para recriação
    example_document: Optional[str] = None  # Documento de exemplo — a IA usa para entender organização e formato esperado
    parent_field_id: Optional[str] = None  # ID do metadado pai — a IA formata como item filho da árvore organizacional
    parent_field_title: Optional[str] = None  # Título do metadado pai


def _parse_spec_and_skill(combined: Optional[str]) -> tuple[str, str]:
    """
    Extrai Spec e Skill de spec_guidelines (formato: --- SPEC ---\\n...\\n--- SKILL ---\\n...).
    Retorna (spec_content, skill_content). O Skill define o papel/tarefa da IA.
    """
    spec_part = ""
    skill_part = ""
    if combined and combined.strip():
        if "--- SKILL ---" in combined:
            before, after = combined.split("--- SKILL ---", 1)
            spec_part = before.replace("--- SPEC ---", "").strip()
            skill_part = after.strip()
        else:
            spec_part = combined.replace("--- SPEC ---", "").strip()
    return spec_part, skill_part


class DocumentSectionForChat(BaseModel):
    """Seção do documento para o chat entender o contexto e ações disponíveis."""
    id: str
    title: str
    index: int  # 1-based para o usuário ("seção 2" = index 2)


class ChatRequest(BaseModel):
    project_id: str
    message: str
    document_id: Optional[str] = None
    document_sections: Optional[list[DocumentSectionForChat]] = None  # Seções editáveis quando document_id presente


class HealthResponse(BaseModel):
    chroma_ready: bool
    collection: str
    persist_dir: str


class AIConfigRequest(BaseModel):
    provider: str  # "ollama" | "groq"
    groq_api_key: Optional[str] = None


# Config de IA em memória (frontend envia via POST /ai-config)
_ai_config: dict = {"provider": "ollama", "groq_api_key": None}


@app.post("/ai-config")
def set_ai_config(request: AIConfigRequest):
    """Recebe configuração de IA do frontend (provider). Chave Groq vem de GROQ_API_KEY no ambiente."""
    global _ai_config
    provider = request.provider if request.provider in ("ollama", "groq") else "ollama"
    _ai_config = {
        "provider": provider,
        "groq_api_key": None,  # Sempre usa GROQ_API_KEY do ambiente (Vercel, etc.)
    }
    return {"status": "ok", "provider": provider}


@app.get("/health", response_model=HealthResponse)
def health():
    """Verifica se ChromaDB está configurado e a collection existe."""
    try:
        col = get_or_create_collection()
        return HealthResponse(
            chroma_ready=True,
            collection=COLLECTION_NAME,
            persist_dir=str(CHROMA_PERSIST_DIR),
        )
    except Exception as e:
        raise HTTPException(status_code=503, detail=str(e))


@app.post("/index")
def index(request: IndexRequest):
    """
    Indexa um documento no ChromaDB.
    - Baixa o arquivo via file_url ou Supabase Storage (file_path)
    - Extrai texto, divide em chunks, gera embeddings
    - Valida duplicação por hash (skip se já indexado, exceto reindex=True)
    - Atualiza status no banco relacional
    """
    try:
        if request.reindex:
            delete_document_from_chroma(request.document_id)

        def check_dup(project_id: str, file_hash: str, current_id: str) -> bool:
            if not get_supabase():
                return False
            r = (
                get_supabase()
                .table("project_materials")
                .select("id")
                .eq("project_id", project_id)
                .eq("file_hash", file_hash)
                .neq("id", current_id)
                .execute()
            )
            return bool(r.data and len(r.data) > 0)

        chunk_count, file_hash = index_document(
            document_id=request.document_id,
            file_url=request.file_url,
            file_name=request.file_name,
            project_id=request.project_id,
            file_path=request.file_path,
            check_duplicate=check_dup if not request.reindex else None,
        )

        supabase = get_supabase()
        if supabase:
            supabase.table("project_materials").update(
                {
                    "status": "PROCESSED",
                    "chunk_count": chunk_count,
                    "file_hash": file_hash,
                }
            ).eq("id", request.document_id).execute()

        return {
            "document_id": request.document_id,
            "chunk_count": chunk_count,
            "file_hash": file_hash,
            "status": "processed",
        }
    except Exception as e:
        supabase = get_supabase()
        if supabase:
            try:
                supabase.table("project_materials").update(
                    {"status": "ERROR"}
                ).eq("id", request.document_id).execute()
            except Exception:
                pass
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/index-positive-feedback")
def index_positive_feedback_endpoint(request: IndexPositiveFeedbackRequest):
    """
    Indexa documento avaliado como 'bom' pelo criador no RAG.
    As seções são adicionadas como exemplos positivos para refinar gerações futuras.
    """
    try:
        sections = [{"title": s.title, "content": s.content} for s in request.sections]
        count = _index_positive_feedback(
            document_id=request.document_id,
            project_id=request.project_id,
            sections=sections,
        )
        return {"status": "ok", "chunks_indexed": count}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/remove-positive-feedback")
def remove_positive_feedback_endpoint(request: DeleteRequest):
    """Remove feedback positivo de um documento (ex.: quando avaliado como 'não ficou bom')."""
    try:
        count = _delete_positive_feedback(request.document_id)
        return {"status": "ok", "chunks_removed": count}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/delete")
def delete_from_rag(request: DeleteRequest):
    """Remove documento do ChromaDB (e opcionalmente atualiza banco)."""
    try:
        count = delete_document_from_chroma(request.document_id)
        supabase = get_supabase()
        if supabase:
            supabase.table("project_materials").update(
                {"chunk_count": 0, "file_hash": None}
            ).eq("id", request.document_id).execute()
        return {"document_id": request.document_id, "chunks_removed": count}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/clear-all")
def clear_all_rag():
    """
    Limpa todo o ChromaDB (base vetorial RAG).
    Mantém regras e orientações (Spec) que estão no código.
    Útil para resetar o sistema quando há erros acumulados.
    """
    try:
        client = get_or_create_client()
        try:
            client.delete_collection(COLLECTION_NAME)
        except Exception:
            pass  # collection pode não existir
        get_or_create_collection()  # recria vazia
        return {"status": "cleared", "message": "Base RAG limpa. Será necessário reindexar os documentos."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/summary")
def get_summary(project_id: str):
    """
    Retorna um resumo do entendimento sobre a documentação da base de dados (RAG).
    Busca chunks do projeto no ChromaDB e gera resumo via Ollama.
    """
    import httpx
    from config import OLLAMA_URL, OLLAMA_SUMMARY_MODEL

    try:
        col = get_or_create_collection()
        # Busca até 30 chunks do projeto no ChromaDB
        results = col.get(
            where={"project_id": {"$eq": project_id}},
            limit=30,
            include=["documents", "metadatas"],
        )
        docs = results.get("documents", [])
        if isinstance(docs, list) and docs and isinstance(docs[0], list):
            docs = docs[0]
        if not docs:
            return {
                "summary": "Ainda não há documentação indexada na base de conhecimento deste projeto. Adicione arquivos (PDF, DOCX, TXT) na aba **Fonte de Dados** para que eu possa analisá-los e ajudá-lo.",
                "sources_count": 0,
            }

        # Concatena chunks (limita tamanho para não sobrecarregar o modelo)
        context = "\n\n---\n\n".join(docs[:20])[:12000]
        prompt = f"""Com base na documentação técnica abaixo, extraída da base de conhecimento do projeto, escreva um resumo objetivo (2 a 4 parágrafos) do entendimento sobre:
- O que o projeto trata
- Principais requisitos, especificações ou conceitos encontrados
- Pontos relevantes para a elaboração de documentos

Documentação:
{context}

Resumo (em português):"""

        with httpx.Client(timeout=90) as client:
            resp = client.post(
                f"{OLLAMA_URL}/api/generate",
                json={
                    "model": OLLAMA_SUMMARY_MODEL,
                    "prompt": prompt,
                    "stream": False,
                    "options": {"temperature": 0.1, "num_predict": 1000},
                },
            )
            resp.raise_for_status()
            summary = resp.json().get("response", "").strip()

        return {
            "summary": summary or "Não foi possível gerar o resumo. Verifique se o Ollama está rodando e o modelo está disponível.",
            "sources_count": len(docs),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/search")
def search(query: str, project_id: Optional[str] = None, n_results: int = 5):
    """Busca semântica no ChromaDB."""
    try:
        col = get_or_create_collection()
        where = {"project_id": {"$eq": project_id}} if project_id else None
        results = col.query(
            query_texts=[query],
            n_results=n_results,
            where=where,
        )
        return {
            "query": query,
            "documents": results.get("documents", [[]])[0],
            "metadatas": results.get("metadatas", [[]])[0],
            "distances": results.get("distances", [[]])[0],
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def _get_rag_context(project_id: str, limit: int = 15) -> str:
    """Obtém contexto da base de conhecimento do projeto."""
    col = get_or_create_collection()
    results = col.get(
        where={"project_id": {"$eq": project_id}},
        limit=limit,
        include=["documents"],
    )
    docs = results.get("documents", [])
    if isinstance(docs, list) and docs and isinstance(docs[0], list):
        docs = docs[0]
    if not docs:
        return ""
    return "\n\n---\n\n".join(docs[:15])[:8000]


def _call_ollama(prompt: str, temperature: float = 0.1, num_predict: int = 800, timeout: int = 240) -> str:
    """Chama Ollama para geração de texto. Retry em caso de Broken pipe ou erro de conexão."""
    import time
    import httpx
    from config import OLLAMA_URL, OLLAMA_SUMMARY_MODEL

    payload = {
        "model": OLLAMA_SUMMARY_MODEL,
        "prompt": prompt,
        "stream": False,
        "options": {"temperature": temperature, "num_predict": num_predict},
    }
    url = f"{OLLAMA_URL}/api/generate"
    last_err = None
    for attempt in range(3):
        try:
            with httpx.Client(timeout=timeout) as client:
                resp = client.post(url, json=payload)
                resp.raise_for_status()
                return resp.json().get("response", "").strip()
        except (BrokenPipeError, ConnectionError, httpx.ConnectError, httpx.ReadTimeout) as e:
            last_err = e
            if attempt < 2:
                time.sleep(1.0 * (attempt + 1))
            else:
                raise
        except OSError as e:
            if getattr(e, "errno", None) == 32:  # Broken pipe
                last_err = e
                if attempt < 2:
                    time.sleep(1.0 * (attempt + 1))
                else:
                    raise
            else:
                raise
    return ""


def _call_groq(prompt: str, temperature: float = 0.1, max_tokens: int = 800, api_key: str = "") -> str:
    """Chama Groq API para geração de texto."""
    from groq import Groq
    from config import GROQ_MODEL

    client = Groq(api_key=api_key)
    completion = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
        max_tokens=max_tokens,
    )
    return (completion.choices[0].message.content or "").strip()


def _call_llm(prompt: str, temperature: float = 0.1, num_predict: int = 800, timeout: int = 240) -> str:
    """
    Chama o provedor de IA configurado (Ollama ou Groq).
    Se Ollama estiver configurado mas indisponível, usa Groq automaticamente (quando chave disponível).
    """
    import os
    from config import GROQ_API_KEY

    provider = _ai_config.get("provider", "ollama")
    groq_key = _ai_config.get("groq_api_key") or os.getenv("GROQ_API_KEY") or GROQ_API_KEY

    # Provedor Groq: usar diretamente
    if provider == "groq" and groq_key:
        return _call_groq(prompt, temperature=temperature, max_tokens=num_predict, api_key=groq_key)

    # Provedor Ollama: tentar Ollama, fallback para Groq se indisponível
    if provider == "ollama":
        try:
            return _call_ollama(prompt, temperature=temperature, num_predict=num_predict, timeout=timeout)
        except Exception as ollama_err:
            err_msg = str(ollama_err).lower()
            if groq_key and (
                "connection" in err_msg or "timeout" in err_msg or "broken pipe" in err_msg or "errno 32" in err_msg
            ):
                return _call_groq(prompt, temperature=temperature, max_tokens=num_predict, api_key=groq_key)
            raise

    # Fallback: tentar Ollama primeiro
    try:
        return _call_ollama(prompt, temperature=temperature, num_predict=num_predict, timeout=timeout)
    except Exception as ollama_err:
        err_msg = str(ollama_err).lower()
        if groq_key and (
            "connection" in err_msg or "timeout" in err_msg or "broken pipe" in err_msg or "errno 32" in err_msg
        ):
            return _call_groq(prompt, temperature=temperature, max_tokens=num_predict, api_key=groq_key)
        raise


@app.post("/generate-document-understanding")
def generate_document_understanding(request: GenerateUnderstandingRequest):
    """
    Lê o documento (estrutura de seções) e a base de conhecimento do projeto.
    Retorna um resumo do entendimento para o usuário confirmar antes de gerar os parágrafos.
    """
    try:
        rag_context = _get_rag_context(request.project_id)
        sections_desc = "\n".join(
            f"- {s.title}" + (f" ({s.helpText})" if s.helpText else "")
            for s in request.sections
        )
        if not rag_context:
            return {
                "summary": "Não há documentação indexada na base de conhecimento deste projeto. "
                "Adicione arquivos na aba Fonte de Dados para que eu possa gerar conteúdo com base no contexto.",
                "sources_count": 0,
            }
        prompt = f"""Você é um assistente especializado em documentação técnica. Analise a estrutura do documento e a base de conhecimento do projeto.

Estrutura do documento (seções a preencher):
{sections_desc}

Base de conhecimento do projeto:
{rag_context}

Escreva um resumo objetivo (2 a 4 parágrafos) do seu entendimento sobre:
1. O que o projeto trata e qual o contexto geral
2. Como você interpreta cada seção do documento com base na base de conhecimento
3. O que você pretende escrever em cada parágrafo ao gerar o documento

Resumo (em português):"""
        summary = _call_llm(prompt, temperature=0.1, num_predict=1000)
        col = get_or_create_collection()
        results = col.get(
            where={"project_id": {"$eq": request.project_id}},
            limit=30,
            include=["documents"],
        )
        docs = results.get("documents", [])
        if isinstance(docs, list) and docs and isinstance(docs[0], list):
            docs = docs[0]
        return {
            "summary": summary or "Não foi possível gerar o resumo. Verifique se o Ollama está rodando.",
            "sources_count": len(docs),
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/generate-section-content")
def generate_section_content(request: GenerateSectionRequest):
    """
    Gera o conteúdo HTML de uma seção com base na base de conhecimento e no Spec do modelo.
    """
    try:
        col = get_or_create_collection()
        query_text = f"{request.section_title} {request.help_text or ''}".strip()
        try:
            results = col.query(
                query_texts=[query_text or "documentação do projeto"],
                n_results=12,
                where={"project_id": {"$eq": str(request.project_id)}},
            )
            raw_docs = results.get("documents") or [[]]
            docs = raw_docs[0] if isinstance(raw_docs, list) and raw_docs else []
            if not isinstance(docs, list):
                docs = []
            docs = [str(d) for d in docs if d]
            rag_context = "\n\n".join(docs[:8])[:6000] if docs else _get_rag_context(request.project_id, limit=12)[:6000]
        except Exception:
            try:
                rag_context = _get_rag_context(request.project_id, limit=12)[:6000]
            except Exception:
                rag_context = ""

        # Posição da seção no documento — CRÍTICO para evitar repetição
        section_idx = request.section_index if request.section_index is not None else 0
        total = request.total_sections if request.total_sections and request.total_sections > 0 else 1
        is_first_section = section_idx == 0

        section_pos_block = f"""
POSIÇÃO NO DOCUMENTO: Esta é a seção {section_idx + 1} de {total} de UM ÚNICO DOCUMENTO.
As seções são PARTES do mesmo documento, NÃO documentos separados.
"""
        if is_first_section:
            section_pos_block += """
PRIMEIRA SEÇÃO: Pode incluir introdução/contexto inicial se pertinente. As seções seguintes NÃO repetirão isso.
"""
        else:
            section_pos_block += """
SEÇÃO INTERMEDIÁRIA/FINAL: É PROIBIDO incluir introdução, objetivos, visão geral, contexto do projeto ou apresentação.
Esses elementos JÁ ESTÃO nas seções anteriores. Comece DIRETO no conteúdo específico desta seção.
NÃO repita cabeçalhos de documento, resumos executivos ou "Este documento apresenta...".
"""

        # Seção filha na árvore organizacional — formatação hierárquica
        parent_hierarchy_block = ""
        if request.parent_field_title and request.parent_field_title.strip():
            parent_hierarchy_block = f"""
HIERARQUIA: Esta seção é FILHA de "{request.parent_field_title.strip()}" na árvore organizacional.
Formate o conteúdo para que o leitor perceba claramente a hierarquia:
- Use listas com marcadores (círculo vazio, <ul><li>) para itens subordinados.
- Cada item deve ser um <li> dentro de <ul>.
- Exemplo: <ul><li>Número (Número automático)</li><li>Título</li><li>Ementa</li></ul>
- O texto deve parecer uma lista indentada sob o pai, como em requisitos funcionais (RF09: ... com subitens).
"""

        # Contexto das seções anteriores: prioriza HTML completo (partes do mesmo documento)
        is_correction_mode = bool(request.creator_feedback and request.creator_feedback.strip())
        ctx_prev = ""
        if request.previous_sections_html and request.previous_sections_html.strip():
            if is_correction_mode:
                # Em modo correção, o documento vai no feedback_block (após "A seguir está o documento original para revisão:")
                ctx_prev = ""
            else:
                ctx_prev = f"""
=== SEÇÕES JÁ ESCRITAS (partes do MESMO documento — NÃO repita) ===
O texto abaixo já foi escrito. Sua seção é CONTINUAÇÃO. PROIBIDO repetir introdução, objetivos, visão geral ou contexto.
---
{request.previous_sections_html[:6000]}
---
=== FIM DAS SEÇÕES ANTERIORES ===
"""
        elif request.document_context:
            ctx_prev = f"\nSeções anteriores (não repita):\n{request.document_context[:800]}\n"

        # Avaliação anterior do criador (quando recriando o documento) — aciona prompt de correção
        feedback_block = ""
        if request.creator_feedback and request.creator_feedback.strip():
            doc_original = ""
            if request.previous_sections_html and request.previous_sections_html.strip():
                doc_original = f"\n{request.previous_sections_html[:6000]}\n"
            feedback_block = f"""
=== PROMPT DE CORREÇÃO (usuário solicitou recriar — documento não ficou bom) ===
{_PROMPT_CORRECAO_DOCUMENTO}
{doc_original}
=== AVALIAÇÃO DO CRIADOR ===
{request.creator_feedback.strip()[:800]}
=== FIM DA AVALIAÇÃO ===

Aplique o prompt de correção acima ao gerar esta seção. Ajuste o conteúdo com base na avaliação para melhorar o resultado.
"""

        # Documento de exemplo: referência de organização e formato esperado
        example_block = ""
        if request.example_document and request.example_document.strip():
            example_excerpt = request.example_document.strip()[:4000]
            example_block = f"""
=== DOCUMENTO DE EXEMPLO (referência de organização e formato — NÃO copiar literalmente) ===
Use como referência para entender: estrutura de seções, tom, formatação, nível de detalhe.
O CONTEÚDO vem dos [DADOS DO PROJETO], mas a organização e estilo devem se assemelhar ao exemplo.
---
{example_excerpt}
---
=== FIM DO EXEMPLO ===

"""

        # Regras constitucionais: Spec/REGRAS_CONSTITUCIONAIS.md (fonte única)
        constitutional_rules = _load_constitutional_rules()

        part_instruction = """REGRAS DE CONTINUIDADE:
- Cada seção é UMA PARTE de um documento maior. NUNCA trate como documento independente.
- NÃO repita introduções, objetivos, visão geral ou contexto entre seções.
- Escreva APENAS o conteúdo específico desta seção, como continuação natural.
- CADA SEÇÃO DEVE TERMINAR COMPLETA: não deixe frases ou parágrafos incompletos. Se o limite de tokens for atingido, encerre a última frase de forma coerente antes de parar.
- Mantenha coerência e estrutura: o documento final em modo de visualização deve fluir como um texto único e contínuo."""

        spec_content, skill_content = _parse_spec_and_skill(request.spec_guidelines)
        # O Skill define o papel e a tarefa da IA. Se não houver Skill, usa fallback.
        task_definition = skill_content[:2000].strip() if skill_content else "Você é um analista de requisitos."

        if spec_content or skill_content:
            spec_excerpt = spec_content[:2500] if spec_content else ""
            doc_type_line = f"Tipo do documento: {request.document_type.strip()}." if request.document_type and request.document_type.strip() else ""
            help_line = f"Orientação da seção (NÃO copie como texto): {request.help_text}" if request.help_text else ""
            prompt = f"""=== REGRAS ABSOLUTAS (violação invalida o documento) ===
{constitutional_rules}
Regra adicional — spec é referência de estrutura: O bloco [ESTRUTURA] abaixo contém EXEMPLOS, INSTRUÇÕES e EXPLICAÇÕES para orientar a IA. Nenhum texto do Spec pode aparecer no documento. Exemplos proibidos de cópia literal: "Objetivo, público-alvo", "O que é o sistema", "Como [perfil], quero [ação]", "[Épico 1]", "[Nome da Feature]", "Incluído / Fora de escopo". Use o Spec APENAS para entender a hierarquia e numeração (EP01 → FT01 → HU001 → RF01). O CONTEÚDO vem exclusivamente dos [DADOS DO PROJETO].
=== FIM DAS REGRAS ===

=== TAREFA (definida pelo documento Skill) ===
{task_definition}

Gere o conteúdo HTML da seção "{request.section_title}".
{section_pos_block}
{parent_hierarchy_block}
{doc_type_line}
{help_line}
{part_instruction}
=== FIM DA TAREFA ===

=== ESTRUTURA (somente hierarquia e numeração — NÃO copiar nenhum texto) ===
{spec_excerpt}
=== FIM DA ESTRUTURA ===
{example_block}
=== DADOS DO PROJETO (única fonte de conteúdo) ===
{rag_context}
=== FIM DOS DADOS ===
{ctx_prev}
{feedback_block}
=== SAÍDA ===
Retorne APENAS HTML válido desta seção (sem <html>, <head>, <body>). Tags permitidas: <h2>, <h3>, <p>, <strong>, <ul>, <li>, <table>, <tr>, <td>, <th>. Proibido: comentários, explicações, raciocínio da IA, texto do Spec, texto copiado da base.

HTML:"""
        else:
            doc_type_line = f"Tipo do documento: {request.document_type.strip()}." if request.document_type and request.document_type.strip() else ""
            help_line = f"Orientação da seção: {request.help_text}" if request.help_text else ""
            prompt = f"""=== REGRAS ABSOLUTAS (violação invalida o documento) ===
{constitutional_rules}
=== FIM DAS REGRAS ===

=== TAREFA (definida pelo documento Skill) ===
{task_definition}

Gere o conteúdo HTML da seção "{request.section_title}" em português formal.
{section_pos_block}
{parent_hierarchy_block}
{doc_type_line}
{help_line}
{part_instruction}
=== FIM DA TAREFA ===
{example_block}
=== DADOS DO PROJETO (única fonte de conteúdo) ===
{rag_context}
=== FIM DOS DADOS ===
{ctx_prev}
{feedback_block}
=== SAÍDA ===
Retorne APENAS HTML válido desta seção (sem <html>, <head>, <body>). Tags permitidas: <h2>, <h3>, <p>, <strong>, <ul>, <li>. Proibido: comentários, explicações ou raciocínio da IA.

HTML:"""

        print(prompt)
        try:
            content = _call_llm(prompt, temperature=0.1, num_predict=2000, timeout=300)
        except Exception as ollama_err:
            err_msg = str(ollama_err)
            if "Broken pipe" in err_msg or "Errno 32" in err_msg:
                raise HTTPException(
                    status_code=503,
                    detail="Ollama fechou a conexão. Verifique se está rodando (ollama serve) e se o modelo está carregado.",
                )
            if "Connection" in err_msg or "timeout" in err_msg.lower():
                raise HTTPException(
                    status_code=503,
                    detail="Não foi possível conectar ao Ollama. Verifique se está rodando em localhost:11434.",
                )
            raise HTTPException(status_code=500, detail=f"Erro ao gerar conteúdo: {err_msg}")
        return {"content": content or ""}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/chat")
def chat(request: ChatRequest):
    """
    Chat com a IA: entende a mensagem do usuário, busca contexto na base de conhecimento (RAG)
    e responde ou sugere ações. Quando document_id é informado, o prompt prioriza o documento em edição.
    """
    try:
        col = get_or_create_collection()
        where_project = {"project_id": {"$eq": request.project_id}}
        results = col.query(
            query_texts=[request.message],
            n_results=10,
            where=where_project,
        )
        docs = results.get("documents", [[]])[0]
        rag_context = "\n\n".join(docs[:8])[:10000] if docs else _get_rag_context(request.project_id, limit=20)

        if not rag_context:
            rag_context = "Ainda não há documentação indexada na base de conhecimento deste projeto."

        if request.document_id:
            sections_info = ""
            if request.document_sections:
                sections_info = "\nSeções do documento (número = índice para o usuário):\n" + "\n".join(
                    f"  {s.index}. {s.title}" for s in request.document_sections
                )
            system_instruction = f"""Você é um assistente especializado em documentação de projetos.
O usuário está editando um documento. Use o contexto da base de conhecimento para responder.
{sections_info}

REGRA CRÍTICA — DIFERENCIE PERGUNTA DE PEDIDO DE ALTERAÇÃO:
- PERGUNTA (como?, quando?, onde?, o que?, qual?): o usuário quer uma RESPOSTA. Responda com base na base de conhecimento. NÃO inclua [SEADOCS_ACTION]. Apenas responda.
- PEDIDO DE ALTERAÇÃO (corrija, refaça, adicione, altere, melhore, inclua no documento): o usuário quer MODIFICAR o documento. Aí sim inclua [SEADOCS_ACTION].

Exemplo PERGUNTA → só responder: "como o cliente faz reserva?" → responda explicando com base nos dados. Sem bloco.
Exemplo PEDIDO → incluir bloco: "adicione na seção 2 como o cliente faz reserva" → inclua [SEADOCS_ACTION].

Quando for PEDIDO de alteração (não pergunta):
- regenerate_all: repetição em várias seções, "refaça o documento", "todas as seções".
- regenerate_section: usuário citar número de seção (ex: "corrija a seção 2").

Bloco só para pedidos de alteração:
[SEADOCS_ACTION]
{{"type": "regenerate_all", "instruction": "instrução completa"}}
[/SEADOCS_ACTION]
ou
[SEADOCS_ACTION]
{{"type": "regenerate_section", "section_index": N, "instruction": "instrução completa"}}
[/SEADOCS_ACTION]

O campo "instruction" deve preservar TODOS os detalhes que o usuário mencionou (ex: se citar "telefone, WhatsApp e Instagram", inclua isso na instruction, não resuma como "vários canais").

REGRA — fora do escopo: Se NÃO há informação na base, responda: "Não há informações sobre isso na documentação deste projeto." """
        else:
            system_instruction = """Você é um assistente especializado em documentação de projetos.
Responda em português, de forma clara e objetiva.
O contexto abaixo é da base de conhecimento do projeto como um todo. Use-o para responder perguntas sobre o projeto.
Se o usuário pedir para criar um novo documento, sugira usar o botão "Novo Documento".

REGRA OBRIGATÓRIA — fora do escopo: Se a pergunta está fora do escopo do projeto ou NÃO existe informação sobre o assunto na base de conhecimento abaixo, responda APENAS: "Não há informações sobre isso na documentação deste projeto." NÃO sugira criar documentos, usar botões, inventar procedimentos ou dar alternativas."""

        prompt = f"""{system_instruction}

Base de conhecimento:
---
{rag_context}
---

Pergunta do usuário: {request.message}

Resposta:"""

        response = _call_llm(prompt, temperature=0.1, num_predict=1500)

        # Parse ação estruturada do bloco [SEADOCS_ACTION] na resposta
        suggested_action = None
        action_match = re.search(r"\[SEADOCS_ACTION\]\s*(.*?)\s*\[/SEADOCS_ACTION\]", response or "", re.DOTALL | re.IGNORECASE)
        if action_match:
            try:
                action_json = json.loads(action_match.group(1).strip())
                action_type = action_json.get("type")
                instruction = (action_json.get("instruction") or "").strip() or None

                # Fallback: se a mensagem indica problema em TODO o documento, forçar regenerate_all
                msg_lower = request.message.lower()
                whole_doc_keywords = ["todas as seções", "todas as sessões", "em todas as seções", "em todas as sessões", "o documento está repetindo", "documento repetindo", "seções repetindo", "refaça", "recreie", "recrie"]
                if any(kw in msg_lower for kw in whole_doc_keywords) and action_type == "regenerate_section":
                    action_type = "regenerate_all"

                # Se instruction está vazia ou é texto de resposta (não instrução), usa a mensagem do usuário
                bad_instructions = ("executando", "correção solicitada", "solicitada", "em andamento")
                if not instruction or any(b in (instruction or "").lower() for b in bad_instructions):
                    instruction = request.message.strip()[:400]

                if action_type == "regenerate_section":
                    si = action_json.get("section_index")
                    if isinstance(si, (int, float)) and si >= 1:
                        suggested_action = {
                            "type": "regenerate_section",
                            "section_index": int(si),
                            "instruction": instruction,
                        }
                elif action_type == "regenerate_all":
                    suggested_action = {
                        "type": "regenerate_all",
                        "instruction": instruction,
                    }
            except (json.JSONDecodeError, TypeError):
                pass

        # SEMPRE remove o bloco [SEADOCS_ACTION] da resposta antes de enviar ao usuário
        raw = response or ""
        raw = re.sub(r"\[\s*SEADOCS_ACTION\s*\][\s\S]*?\[\s*/\s*SEADOCS_ACTION\s*\]", "", raw, flags=re.IGNORECASE)
        raw = re.sub(r"\[\s*SEADOCS_ACTION\s*\][\s\S]*$", "", raw, flags=re.IGNORECASE)  # fallback se fechamento ausente
        response = re.sub(r"\n{3,}", "\n\n", raw).strip()

        # Fallback: detecta pedidos de geração por keywords (compatibilidade)
        if suggested_action is None:
            msg_lower = request.message.lower().strip()
            gerar_keywords = ["gere", "gerar", "crie", "criar", "preencha", "preencher", "escreva", "escrever", "produza", "produzir"]
            if any(kw in msg_lower for kw in gerar_keywords) and any(w in msg_lower for w in ["documento", "documentos", "parágrafo", "parágrafos", "seção", "seções"]):
                suggested_action = "generate_document"

        return {
            "response": response or "Não foi possível gerar uma resposta. Verifique se o Claude (ANTHROPIC_API_KEY) ou Ollama está configurado.",
            "suggested_action": suggested_action,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class ProposeSectionsRequest(BaseModel):
    """Quando o modelo não tem campos de metadados, a IA analisa a base e propõe seções."""
    project_id: str
    document_type: Optional[str] = None
    spec_guidelines: Optional[str] = None


class PlanDocumentRequest(BaseModel):
    project_id: str
    sections: list[DocumentSectionInput]


@app.post("/propose-sections-from-knowledge-base")
def propose_sections_from_knowledge_base(request: ProposeSectionsRequest):
    """
    Quando o modelo não tem campos de metadados definidos, analisa a base de conhecimento
    do projeto e propõe as seções necessárias para construir o documento completo.
    A IA entende os problemas/necessidades com base nos dados e cria todos os campos
    que precisar para todas as partes do documento.
    """
    rag_context = _get_rag_context(request.project_id, limit=25)
    if not rag_context:
        return {
            "sections": [
                {"id": "section1", "title": "Conteúdo do documento", "helpText": "Adicione arquivos na base de conhecimento do projeto para que a IA possa propor seções automaticamente."}
            ]
        }

    doc_type = (request.document_type or "documento").strip()
    spec_block = ""
    if request.spec_guidelines and request.spec_guidelines.strip():
        spec_block = f"""
ESTRUTURA ESPERADA (Spec do modelo — use como referência de hierarquia e numeração):
---
{request.spec_guidelines[:2000]}
---
"""

    prompt = f"""Você é um analista de documentação. Analise a base de conhecimento do projeto abaixo e identifique TODOS os problemas, necessidades, requisitos e temas que devem ser cobertos em um documento do tipo "{doc_type}".

Base de conhecimento do projeto:
---
{rag_context[:8000]}
---
{spec_block}

Com base na análise, liste as SEÇÕES que o documento deve ter para cobrir completamente o conteúdo relevante. Cada seção deve ser um campo/metadado que a IA preencherá ao gerar o documento.
- Use títulos claros e objetivos para cada seção
- Ordene de forma lógica (introdução/contexto primeiro, detalhes depois)
- Inclua todas as partes necessárias para um documento completo e profissional

Retorne APENAS uma lista no formato abaixo, uma seção por linha:
id:título da seção

Exemplo:
sec-1:Introdução e contexto do projeto
sec-2:Objetivos e escopo
sec-3:Requisitos funcionais identificados
sec-4:Requisitos não funcionais
sec-5:Regras de negócio
sec-6:Considerações finais

Lista de seções:"""

    try:
        response = _call_llm(prompt, temperature=0.1, num_predict=1200, timeout=180)
    except Exception:
        response = ""

    sections: list[dict] = []
    seen_ids: set[str] = set()
    for i, line in enumerate(response.splitlines()):
        line = line.strip()
        if not line or ":" not in line:
            continue
        try:
            idx = line.index(":")
            sec_id = line[:idx].strip().replace(" ", "-") or f"sec-{i}"
            sec_id = sec_id[:50]
            if sec_id in seen_ids:
                suffix = 0
                while f"{sec_id}-{suffix}" in seen_ids:
                    suffix += 1
                sec_id = f"{sec_id}-{suffix}"
            seen_ids.add(sec_id)
            title = line[idx + 1 :].strip()
            if title and len(title) > 2:
                sections.append({
                    "id": sec_id,
                    "title": title[:200],
                    "helpText": f'Conteúdo de "{title}" com base na base de conhecimento.',
                })
        except ValueError:
            continue

    if not sections:
        sections = [
            {"id": "sec-1", "title": "Conteúdo principal", "helpText": "Conteúdo do documento com base na base de conhecimento."}
        ]

    return {"sections": sections}


@app.post("/plan-document-structure")
def plan_document_structure(request: PlanDocumentRequest):
    """
    Fase de planejamento da geração de documentos com seções repetíveis.

    Para cada seção marcada como 'repeatable', consulta a base de conhecimento do projeto
    e identifica quantas instâncias são necessárias (ex: quantos épicos, features, histórias).

    Retorna uma lista expandida e ordenada de seções prontas para geração, mantendo
    a ordem do template e substituindo cada seção repetível pelas suas instâncias concretas.
    """
    rag_context = _get_rag_context(request.project_id, limit=20)
    repeatable = [s for s in request.sections if s.repeatable]

    # Sem seções repetíveis ou sem RAG → retorna o template como está
    if not repeatable or not rag_context:
        return {
            "planned_sections": [
                {
                    "id": s.id,
                    "title": s.title,
                    "helpText": s.helpText or f'Conteúdo de "{s.title}"',
                    "templateSectionId": s.id,
                }
                for s in request.sections
            ]
        }

    # Monta prompt para o LLM identificar instâncias de cada seção repetível
    categories_desc = "\n".join(
        f'[{s.id}] {s.title}: {s.planningInstruction or "Identifique todos os itens deste tipo encontrados no projeto"}'
        for s in repeatable
    )

    prompt = f"""Você é um analista de requisitos. Analise a documentação do projeto abaixo e identifique todos os elementos de cada categoria.

Documentação do projeto:
{rag_context[:4000]}

Categorias a identificar:
{categories_desc}

Para cada categoria, liste os títulos específicos encontrados no projeto.
Use EXATAMENTE o formato abaixo, uma linha por item encontrado:
[id_categoria]: Título específico do item

Exemplo:
[epic-0]: Autenticação e Controle de Acesso
[epic-0]: Gestão Documental
[feature-1]: Login com Gov.br
[feature-1]: Upload e Versionamento de Documentos
[story-2]: Como usuário, quero me autenticar via Gov.br para acessar o sistema

Lista completa:"""

    try:
        response = _call_llm(prompt, temperature=0.1, num_predict=1500, timeout=240)
    except Exception:
        response = ""

    # Faz parse das linhas de resposta
    instances_map: dict[str, list[str]] = {s.id: [] for s in repeatable}
    for line in response.splitlines():
        line = line.strip()
        if not line or not line.startswith("["):
            continue
        try:
            bracket_end = line.index("]:")
            sec_id = line[1:bracket_end].strip()
            title = line[bracket_end + 2:].strip()
            if sec_id in instances_map and title:
                instances_map[sec_id].append(title)
        except ValueError:
            continue

    # Fallback: se o LLM não retornou nada para uma seção, mantém uma instância
    for s in repeatable:
        if not instances_map[s.id]:
            instances_map[s.id] = [s.title]

    # Expande o template mantendo a ordem original
    planned: list[dict] = []
    for s in request.sections:
        if not s.repeatable:
            planned.append({
                "id": s.id,
                "title": s.title,
                "helpText": s.helpText or f'Conteúdo de "{s.title}"',
                "templateSectionId": s.id,
            })
        else:
            for i, instance_title in enumerate(instances_map[s.id][:30]):  # cap 30 instâncias
                planned.append({
                    "id": f"{s.id}-inst-{i}",
                    "title": instance_title,
                    "helpText": f'{s.helpText or s.title}: {instance_title}',
                    "templateSectionId": s.id,
                })

    return {"planned_sections": planned}


class AnalyzeTemplateRequest(BaseModel):
    file_name: str


@app.get("/list-template-examples")
def list_template_examples():
    """Lista arquivos DOCX disponíveis em public/modelos/ para geração automática de modelos."""
    if not MODELOS_DIR.exists():
        return {"files": []}
    files = sorted(f.name for f in MODELOS_DIR.glob("*.docx"))
    return {"files": files}


def _rgb_to_hex(rgb_color) -> Optional[str]:
    try:
        if rgb_color is None:
            return None
        return f"#{rgb_color.red:02X}{rgb_color.green:02X}{rgb_color.blue:02X}"
    except Exception:
        return None


def _pt(length) -> Optional[float]:
    try:
        if length is None:
            return None
        return round(length.pt, 1)
    except Exception:
        return None


def _align_css(alignment) -> str:
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return "center"
        if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return "right"
        if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return "justify"
    except Exception:
        pass
    return "left"


def _heading_styles(para) -> str:
    """Extrai estilos inline de um parágrafo heading (cor, tamanho, fonte, alinhamento)."""
    styles = []
    align = _align_css(para.alignment)
    if align != "left":
        styles.append(f"text-align:{align}")
    # Pega formatação do primeiro run não-vazio
    for run in para.runs:
        if not run.text.strip():
            continue
        color = _rgb_to_hex(run.font.color.rgb if run.font.color and run.font.color.type else None)
        if color and color.upper() != "#000000":
            styles.append(f"color:{color}")
        size = _pt(run.font.size)
        if size:
            styles.append(f"font-size:{int(size)}pt")
        if run.font.name:
            styles.append(f"font-family:'{run.font.name}'")
        break
    return ";".join(styles)


def _classify_para(para) -> tuple[bool, str, str]:
    """
    Classifica o parágrafo e retorna (é_heading, tag_html, texto_limpo).

    Detecta cabeçalhos de seção em qualquer tipo de documento, independente de estilo formal:
      1. Estilos de Heading (Word/LibreOffice): Heading 1-4, Título, Title
      2. Texto curto todo em maiúsculas (ex: "OBJETO DO CONTRATO")
      3. Parágrafo curto inteiramente em negrito, sem ponto final (ex: "Participantes")
      4. Prefixo numerado arábico com texto curto (ex: "1. Abertura", "2.1 Pauta")
      5. Prefixo numerado romano (ex: "I – Das Disposições", "II. Objeto")
      6. Prefixo de letra maiúscula (ex: "A – Fatos", "B) Direito")
      7. Palavras-chave de cabeçalho típicas de documentos normativos/jurídicos/administrativos
         (ex: "CLÁUSULA", "ARTIGO", "CONSIDERANDO", "DELIBERAÇÃO", "RESOLVE:")

    Parágrafos de corpo NÃO são headings — seu conteúdo não será copiado no template.
    """
    import re

    style_name = (para.style.name or "") if para.style else ""
    text = para.text.strip()

    if not text:
        return False, "p", ""

    # ── 1. Estilo formal de Heading ──────────────────────────────────────────
    heading_tags = {
        "Heading 1": "h1", "Título 1": "h1", "Title": "h1", "Título": "h1",
        "Heading 2": "h2", "Título 2": "h2",
        "Heading 3": "h3", "Título 3": "h3",
        "Heading 4": "h4", "Título 4": "h4",
    }
    tag = next((v for k, v in heading_tags.items() if k in style_name), "p")
    if tag != "p":
        return True, tag, text

    # Limite de tamanho: títulos raramente têm mais de 160 caracteres
    if len(text) > 160:
        return False, "p", text

    # ── 2. Texto curto todo em maiúsculas ────────────────────────────────────
    if text.isupper() and len(text) >= 3:
        return True, "h3", text

    # ── 3. Parágrafo inteiramente em negrito e curto (sem ponto final de frase)
    runs_text = "".join(r.text for r in para.runs if r.text.strip())
    all_bold = bool(runs_text) and all(
        r.bold for r in para.runs if r.text.strip()
    )
    if all_bold and len(text) < 120 and not text.endswith("."):
        return True, "h3", text

    # ── 4. Prefixo numérico arábico (1., 2., 1.1, 2.3.1 …) ──────────────────
    if re.match(r"^\d+(\.\d+)*\.?\s+\S", text) and len(text) < 120:
        return True, "h3", text

    # ── 5. Prefixo romano (I –, II., III) ────────────────────────────────────
    if re.match(r"^(I{1,3}|IV|V|VI{0,3}|IX|X{1,3})\s*[\.\-–—]\s*\S", text, re.IGNORECASE) and len(text) < 120:
        return True, "h3", text

    # ── 6. Prefixo letra maiúscula (A –, B), C. …) ───────────────────────────
    if re.match(r"^[A-Z]\s*[\.\-–—\)]\s*\S", text) and len(text) < 120:
        return True, "h3", text

    # ── 7. Palavras-chave de cabeçalho normativo/jurídico/administrativo ─────
    normative_keywords = [
        r"^cláusula\s+\w", r"^artigo\s+\d", r"^art\.\s*\d",
        r"^parágrafo\s+\w", r"^§\s*\d",
        r"^considerando\b", r"^resolve\s*:", r"^delibera\s*:",
        r"^decreta\s*:", r"^portaria\b", r"^resolução\b",
        r"^ementa\s*:", r"^objeto\s*:", r"^vigência\s*:",
        r"^das?\s+disposiç", r"^das?\s+obrigaç",
        r"^seção\s+\w", r"^capítulo\s+\w",
    ]
    lower = text.lower()
    for kw in normative_keywords:
        if re.match(kw, lower):
            return True, "h3", text

    return False, "p", text


def _unique_cells(row):
    """Retorna células únicas de uma linha (descarta repetições de mesclagem)."""
    seen = set()
    result = []
    for cell in row.cells:
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            result.append(cell)
    return result


def _table_to_html(table) -> str:
    """
    Converte tabela docx em HTML preservando estrutura (linhas × colunas + estilo de borda).

    Regra de conteúdo:
    - Primeira linha: considerada cabeçalho estrutural → mantém o TEXTO (é um rótulo,
      não um dado). Ex: "Nome", "Data", "Responsável" são parte do layout da tabela.
    - Demais linhas: células ficam vazias — o conteúdo será gerado pela IA no projeto.
    """
    rows = table.rows
    if not rows:
        return ""

    rows_html = ""
    for row_idx, row in enumerate(rows):
        cells = _unique_cells(row)
        cells_html = ""
        td_style = "border:1px solid #ccc;padding:6px 10px;vertical-align:top"

        for col_idx, cell in enumerate(cells):
            if row_idx == 0:
                # Cabeçalho: mantém o rótulo real da coluna (estrutural, não é dado do projeto)
                header_text = cell.text.strip()
                if not header_text:
                    header_text = f"Coluna {col_idx + 1}"
                safe = header_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                cells_html += (
                    f'<th style="{td_style};background:#f3f4f6;font-weight:600">{safe}</th>'
                )
            else:
                cells_html += f'<td style="{td_style}"></td>'

        rows_html += f"<tr>{cells_html}</tr>"

    return (
        '<table style="border-collapse:collapse;width:100%;margin:12px 0">'
        f"<tbody>{rows_html}</tbody></table>"
    )


# Padrões que indicam seções naturalmente repetíveis em qualquer tipo de documento.
# Cada entrada: (palavras-chave no título, instrução de planejamento, nome genérico do campo)
# O nome genérico substitui o título específico do exemplo no template gerado,
# pois o modelo deve ser reutilizável em projetos diferentes.
_REPEATABLE_PATTERNS: list[tuple[list[str], str, str]] = [
    # ── Engenharia de Software / Requisitos ──────────────────────────────────
    (
        ["épico", "epic", "epico"],
        "Liste todos os épicos/módulos/subsistemas identificados na documentação do projeto",
        "Épico",
    ),
    (
        ["feature", "funcionalidade"],
        "Liste todas as features/funcionalidades identificadas na documentação do projeto",
        "Feature",
    ),
    (
        ["história de usuário", "historia de usuario", "user story"],
        "Liste todas as histórias de usuário identificadas na documentação do projeto",
        "História de Usuário",
    ),
    (
        ["caso de uso", "use case"],
        "Liste todos os casos de uso identificados na documentação do projeto",
        "Caso de Uso",
    ),
    (
        ["requisito funcional"],
        "Liste todos os requisitos funcionais identificados na documentação do projeto",
        "Requisito Funcional",
    ),
    (
        ["requisito não funcional", "requisito nao funcional", "requisito de qualidade"],
        "Liste todos os requisitos não funcionais identificados na documentação do projeto",
        "Requisito Não Funcional",
    ),
    (
        ["sprint", "iteração", "iteracao"],
        "Liste todos os sprints/iterações planejados para o projeto",
        "Sprint",
    ),
    (
        ["risco", "risk"],
        "Liste todos os riscos identificados na documentação do projeto",
        "Risco",
    ),
    (
        ["ator", "stakeholder", "parte interessada"],
        "Liste todos os atores/stakeholders identificados na documentação do projeto",
        "Ator / Stakeholder",
    ),
    # ── Documentos Jurídicos / Normativos ────────────────────────────────────
    (
        ["cláusula", "clausula"],
        "Liste todas as cláusulas necessárias para este tipo de documento",
        "Cláusula",
    ),
    (
        ["artigo", "art."],
        "Liste todos os artigos necessários para este documento normativo",
        "Artigo",
    ),
    (
        ["considerando"],
        "Liste todos os considerandos necessários para este documento",
        "Considerando",
    ),
    (
        ["parágrafo único", "paragrafo unico"],
        "Liste os parágrafos únicos ou complementares necessários",
        "Parágrafo",
    ),
    # ── Atas / Minutas ───────────────────────────────────────────────────────
    (
        ["ponto de pauta", "item de pauta", "assunto tratado"],
        "Liste todos os pontos de pauta / assuntos a serem tratados na reunião",
        "Ponto de Pauta",
    ),
    (
        ["deliberação", "deliberacao", "resolução", "resolucao"],
        "Liste todas as deliberações/resoluções a serem registradas",
        "Deliberação",
    ),
    (
        ["encaminhamento", "ação", "acao", "providência"],
        "Liste todos os encaminhamentos/ações definidos na reunião",
        "Encaminhamento",
    ),
    # ── Recursos / Defesas ───────────────────────────────────────────────────
    (
        ["fundamento", "argumento", "razão", "razao"],
        "Liste os fundamentos/argumentos relevantes para este recurso ou defesa",
        "Fundamento",
    ),
    # ── Módulo / Componente (genérico) ───────────────────────────────────────
    (
        ["módulo", "modulo", "componente", "subsistema"],
        "Liste todos os módulos/componentes identificados na documentação do projeto",
        "Módulo",
    ),
]

# Prefixos numerados típicos de documentos de requisitos (EP01, FT02, HU03, etc.)
# mapeados para o nome genérico correspondente
_NUMBERED_PREFIX_MAP: list[tuple[str, str]] = [
    ("ep",  "Épico"),
    ("ft",  "Feature"),
    ("hu",  "História de Usuário"),
    ("us",  "História de Usuário"),
    ("uc",  "Caso de Uso"),
    ("rf",  "Requisito Funcional"),
    ("rnf", "Requisito Não Funcional"),
    ("sp",  "Sprint"),
]


def _detect_repeatable(title: str) -> tuple[bool, Optional[str], Optional[str]]:
    """
    Verifica se o título indica uma seção naturalmente repetível.
    Retorna (é_repetível, instrução_planejamento, nome_genérico).

    Detecta dois padrões:
      1. Prefixo numerado: EP01, FT02, HU03, US04, UC05, RF06, RNF07, SP08, R09...
      2. Palavras-chave no título: "épico", "feature", "história de usuário"...
    """
    import re
    lower = title.lower().strip()

    # 1. Prefixo numerado (ex: "EP01 – Gestão Documental", "FT-03: Upload")
    for prefix, generic_name in _NUMBERED_PREFIX_MAP:
        pattern = rf"^{re.escape(prefix)}\s*[\-–—]?\s*\d"
        if re.match(pattern, lower):
            # Localiza a instrução de planejamento pelo nome genérico
            for keywords, instruction, gn in _REPEATABLE_PATTERNS:
                if gn == generic_name:
                    return True, instruction, generic_name
            return True, f"Liste todos os itens do tipo {generic_name} encontrados no projeto", generic_name

    # 2. Palavra-chave no título
    for keywords, instruction, generic_name in _REPEATABLE_PATTERNS:
        if any(kw in lower for kw in keywords):
            return True, instruction, generic_name

    return False, None, None


def _metadata_field_html(
    field_id: str,
    title: str,
    help_text: str,
    repeatable: bool = False,
    planning_instruction: Optional[str] = None,
) -> str:
    """Gera HTML de um blot MetadataField compatível com o Quill."""
    safe_title = title.replace('"', "&quot;").replace("<", "&lt;").replace(">", "&gt;")
    safe_help = help_text.replace('"', "&quot;").replace("<", "&lt;").replace(">", "&gt;")
    safe_instruction = (planning_instruction or "").replace('"', "&quot;").replace("<", "&lt;").replace(">", "&gt;")
    repeatable_attr = ' data-repeatable="true"' if repeatable else ''
    instruction_attr = f' data-planning-instruction="{safe_instruction}"' if safe_instruction else ''
    repeatable_badge = (
        '<span style="font-size:9px;background:#dbeafe;color:#1d4ed8;'
        'padding:1px 6px;border-radius:99px;margin-left:6px;font-weight:600">REPETÍVEL</span>'
    ) if repeatable else ''
    return (
        f'<div class="sgid-metadata-field" contenteditable="false" '
        f'data-field-id="{field_id}" data-field-title="{safe_title}" '
        f'data-field-help="{safe_help}" data-topic-id=""'
        f'{repeatable_attr}{instruction_attr}>'
        f'<div class="sgid-metadata-field__header">'
        f'<div class="sgid-metadata-field__title">{safe_title}{repeatable_badge}</div>'
        f'<div class="sgid-metadata-field__help" style="font-size:10px;opacity:0.7">{safe_help}</div>'
        f"</div>"
        f'<div class="sgid-metadata-field__textarea">Campo dinâmico — a IA criará as instâncias necessárias ao gerar o documento.</div>'
        f"</div>"
    ) if repeatable else (
        f'<div class="sgid-metadata-field" contenteditable="false" '
        f'data-field-id="{field_id}" data-field-title="{safe_title}" '
        f'data-field-help="{safe_help}" data-topic-id="">'
        f'<div class="sgid-metadata-field__header">'
        f'<div class="sgid-metadata-field__title">{safe_title}</div>'
        f'<div class="sgid-metadata-field__help" style="font-size:10px;opacity:0.7">{safe_help}</div>'
        f"</div>"
        f'<div class="sgid-metadata-field__textarea">Digite aqui (campo editável no documento)...</div>'
        f"</div>"
    )


@app.post("/analyze-document-template")
def analyze_document_template(request: AnalyzeTemplateRequest):
    """
    Analisa um arquivo DOCX e extrai APENAS a estrutura para criar um modelo de documento.
    O conteúdo dos parágrafos de corpo NÃO é copiado — pertence ao documento de exemplo.
    O que é extraído:
      - Hierarquia de títulos/seções com estilos visuais (cor, tamanho, fonte, alinhamento)
      - Um campo de metadado (blot) vazio após cada seção para preenchimento posterior
      - Estrutura de tabelas (linhas × colunas) sem dados — células ficam vazias
    Retorna:
      - template_html: HTML estrutural com blots de seção vazios
      - sections: lista de seções identificadas (id, title, helpText)
      - suggested_name: nome sugerido para o modelo
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx não instalado. Execute: pip install python-docx")

    file_path = MODELOS_DIR / request.file_name
    if not file_path.exists():
        raise HTTPException(status_code=404, detail=f"Arquivo não encontrado: {request.file_name}")

    try:
        doc = DocxDocument(str(file_path))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao abrir DOCX: {e}")

    html_parts: list[str] = []
    sections: list[dict] = []
    heading_count = 0
    # Controla quais tipos repetíveis já foram adicionados ao template
    # para evitar duplicatas (EP01, EP02... → apenas UM campo "Épico")
    seen_repeatable_types: set[str] = set()

    # Mapa de elementos body → objetos python-docx para preservar ordem
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    body = doc.element.body
    for child in body:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if local == "p":
            para = para_map.get(child)
            if para is None:
                continue

            is_heading, tag, clean_text = _classify_para(para)

            if not is_heading:
                # Parágrafos de corpo: NÃO copiar conteúdo.
                continue

            if not clean_text:
                continue

            # Detecta se é uma seção repetível e qual é o seu tipo genérico
            repeatable, planning_instruction, generic_name = _detect_repeatable(clean_text)

            if repeatable and generic_name:
                # Deduplica: se já existe um campo para este tipo, ignora instâncias adicionais.
                # EP01, EP02, EP03 → colapsa em UM campo genérico "Épico".
                if generic_name in seen_repeatable_types:
                    continue
                seen_repeatable_types.add(generic_name)

                # Usa o nome genérico como título — não o texto específico do exemplo.
                # O modelo é reutilizável; os nomes reais serão gerados pela IA no projeto.
                section_title = generic_name
                inline_styles = _heading_styles(para)
                style_attr = f' style="{inline_styles}"' if inline_styles else ""
                safe_label = generic_name.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                html_parts.append(f"<{tag}{style_attr}>{safe_label}</{tag}>")
            else:
                # Seção fixa: mantém título estrutural (ex: "Visão Geral", "Objetivo")
                section_title = clean_text
                inline_styles = _heading_styles(para)
                style_attr = f' style="{inline_styles}"' if inline_styles else ""
                safe_text = clean_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                html_parts.append(f"<{tag}{style_attr}>{safe_text}</{tag}>")

            # Campo de metadado vazio — será preenchido pela IA ao criar o documento
            field_id = f"section-{heading_count}"
            heading_count += 1
            help_text = (
                f'{section_title}: gerado pela IA com base na base de conhecimento do projeto'
                if repeatable
                else f'Conteúdo da seção "{section_title}"'
            )

            sections.append({
                "id": field_id,
                "title": section_title,
                "helpText": help_text,
                "repeatable": repeatable,
                "planningInstruction": planning_instruction,
            })
            html_parts.append(_metadata_field_html(field_id, section_title, help_text, repeatable, planning_instruction))
            html_parts.append("<p><br></p>")

        elif local == "tbl":
            table = table_map.get(child)
            if table is None:
                continue
            # Tabela: preserva estrutura (linhas × colunas + estilo), NÃO copia dados
            html_parts.append(_table_to_html(table))
            html_parts.append("<p><br></p>")

    # Nome sugerido baseado no nome do arquivo sem extensão
    suggested_name = Path(request.file_name).stem

    return {
        "template_html": "\n".join(html_parts),
        "sections": sections,
        "suggested_name": suggested_name,
    }


def run():
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)


if __name__ == "__main__":
    run()
